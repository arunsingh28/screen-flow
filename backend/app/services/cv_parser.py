"""
CV Parser Service
Handles PDF/DOCX parsing and LLM-based CV extraction with skill recency tracking
"""

import json
import io
from typing import Dict, Any, Optional
import PyPDF2
import pdfplumber
from docx import Document
from sqlalchemy.orm import Session
from app.services.bedrock import bedrock_service
from app.models.jd_builder import LLMCallType, CVParseDetail
from app.models.job import CV
import logging

logger = logging.getLogger(__name__)


CV_PARSING_PROMPT = """You are an expert CV parser. Extract structured information from this CV/resume with special attention to skill recency and usage timeline.

CV TEXT:
{cv_text}

Extract ALL information accurately. Don't hallucinate or infer details not present.

Return ONLY this JSON:

{{
  "personal_info": {{
    "name": "extracted or null",
    "email": "extracted or null",
    "phone": "extracted or null",
    "location": "current city/country or null",
    "linkedin": "url or null",
    "github": "url or null",
    "portfolio": "url or null"
  }},

  "summary": {{
    "total_experience_years": calculated_from_work_history,
    "current_role": "most recent job title",
    "current_company": "most recent company",
    "career_level": "inferred: entry/mid/senior/lead based on experience and roles",
    "currently_employed": true/false
  }},

  "work_experience": [
    {{
      "company": "company name",
      "role": "job title",
      "start_date": "YYYY-MM or YYYY",
      "end_date": "YYYY-MM or 'Present'",
      "duration_months": calculated,
      "location": "city/country if mentioned",
      "employment_type": "full-time/part-time/contract/internship if mentioned",
      "company_type": "startup/mid-size/MNC/agency inferred from context or null",
      "industry": "FinTech/HealthTech/E-commerce/SaaS etc. if mentioned or inferred",
      "responsibilities": ["extracted key points"],
      "achievements": ["quantified achievements with numbers/impact"],
      "tech_stack": ["technologies mentioned for this role"],
      "team_size": "number if mentioned",
      "reporting_to": "role if mentioned"
    }}
  ],

  "skills": {{
    "technical_skills": [
      {{
        "skill": "skill name",
        "proficiency": "beginner/intermediate/advanced/expert - inferred from context",
        "total_years_experience": "calculated from all usage across work history",
        "first_used": "YYYY-MM or YYYY - earliest mention in work history",
        "last_used": "YYYY-MM or YYYY or 'Present' - most recent usage",
        "months_since_last_used": calculated_number_or_0_if_current,
        "recency_score": 0-100,
        "skill_freshness": "current/recent/outdated",
        "usage_timeline": [
          {{
            "company": "where used",
            "role": "role title",
            "period": "YYYY-MM to YYYY-MM or Present",
            "duration_months": calculated,
            "usage_context": "brief description of how skill was used",
            "usage_intensity": "primary/secondary/occasional"
          }}
        ],
        "evidence": "summary of where/how this skill was demonstrated across career",
        "skill_depth_indicators": ["led projects", "mentored others", "architectural decisions", "production scale", etc.]
      }}
    ],
    "soft_skills": ["communication", "leadership", etc.],
    "tools": ["specific tools mentioned"],
    "languages": {{
      "programming": [
        {{
          "language": "language name",
          "proficiency": "level if mentioned",
          "last_used": "YYYY-MM or Present"
        }}
      ],
      "spoken": [
        {{
          "language": "language name",
          "proficiency": "native/fluent/intermediate/basic"
        }}
      ]
    }}
  }},

  "education": [
    {{
      "degree": "degree name",
      "field": "field of study",
      "institution": "college/university name",
      "institution_tier": "IIT/NIT/Tier-1/Tier-2/Other - if recognizable",
      "graduation_year": "YYYY or null",
      "start_year": "YYYY or null",
      "grade": "CGPA/percentage if mentioned",
      "relevant_coursework": ["if mentioned"]
    }}
  ],

  "projects": [
    {{
      "name": "project name",
      "description": "brief description",
      "date": "YYYY-MM or YYYY if mentioned",
      "technologies": ["tech used"],
      "impact": "quantified impact if mentioned",
      "scale": "users/requests/data size if mentioned",
      "link": "url if provided",
      "role": "individual/team lead/contributor",
      "is_recent": true/false (within last 2 years)
    }}
  ],

  "certifications": [
    {{
      "name": "certification name",
      "issuer": "issuing organization",
      "date": "YYYY-MM or YYYY",
      "expiry_date": "YYYY-MM or YYYY or null",
      "credential_id": "if mentioned",
      "is_active": true/false (not expired)
    }}
  ],

  "additional": {{
    "notice_period": "extracted if mentioned (immediate/15days/30days/60days/90days)",
    "current_ctc": "extracted if mentioned, in LPA",
    "expected_ctc": "extracted if mentioned, in LPA",
    "willing_to_relocate": true/false/null,
    "open_source": [
      {{
        "project": "project name or url",
        "role": "maintainer/contributor",
        "description": "brief description if available",
        "last_activity": "YYYY-MM or YYYY if mentioned"
      }}
    ],
    "publications": [
      {{
        "title": "paper title",
        "venue": "conference/journal",
        "year": "YYYY",
        "link": "url if provided"
      }}
    ],
    "awards": [
      {{
        "award": "award name",
        "issuer": "organization",
        "year": "YYYY",
        "description": "brief description if available"
      }}
    ],
    "github_stats": {{
      "username": "extracted if mentioned",
      "contributions_mentioned": true/false,
      "active_recently": true/false (if dates mentioned)
    }}
  }},

  "career_trajectory": {{
    "progression": "upward/lateral/downward/mixed - based on role changes",
    "job_changes_count": number,
    "average_tenure_months": calculated,
    "longest_tenure_months": calculated,
    "fastest_promotion": "description if evident",
    "career_gaps": [
      {{
        "after_company": "company name",
        "gap_duration_months": calculated,
        "gap_period": "YYYY-MM to YYYY-MM"
      }}
    ]
  }},

  "meta": {{
    "cv_quality": "high/medium/low - how well-structured and detailed",
    "parsing_confidence": "high/medium/low",
    "cv_format_issues": ["missing dates", "vague descriptions", etc.],

    "red_flags": [
      {{
        "flag": "specific issue",
        "severity": "high/medium/low",
        "details": "explanation"
      }}
    ],

    "standout_points": [
      {{
        "point": "unique achievement or credential",
        "category": "experience/education/project/achievement",
        "relevance": "why this stands out"
      }}
    ],

    "skill_freshness_summary": {{
      "current_skills_count": number,
      "outdated_skills_count": number,
      "actively_learning": ["skills with recent usage or side projects"]
    }}
  }}
}}

CRITICAL RULES:
1. Calculate skill recency based on work history dates
2. Map skills to specific jobs and timeframes
3. Identify red flags (job hopping, gaps, outdated skills)
4. Extract quantified achievements with numbers
5. Be conservative - don't make up information
6. Return ONLY valid JSON, no additional text"""


class CVParserService:
    """Service for parsing CVs and extracting structured data"""

    def extract_text_from_pdf(self, file_content: bytes) -> str:
        """Extract text from PDF file using pdfplumber (better for complex layouts)"""
        try:
            pdf_file = io.BytesIO(file_content)
            text_parts = []

            with pdfplumber.open(pdf_file) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        text_parts.append(text)

            return "\n\n".join(text_parts)

        except Exception as e:
            logger.warning(f"pdfplumber failed, falling back to PyPDF2: {e}")
            # Fallback to PyPDF2
            try:
                pdf_file = io.BytesIO(file_content)
                pdf_reader = PyPDF2.PdfReader(pdf_file)
                text_parts = []

                for page in pdf_reader.pages:
                    text = page.extract_text()
                    if text:
                        text_parts.append(text)

                return "\n\n".join(text_parts)
            except Exception as e2:
                logger.error(f"PDF parsing failed: {e2}")
                raise ValueError(f"Failed to parse PDF: {e2}")

    def extract_text_from_docx(self, file_content: bytes) -> str:
        """Extract text from DOCX file"""
        try:
            docx_file = io.BytesIO(file_content)
            doc = Document(docx_file)
            text_parts = []

            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)

            return "\n".join(text_parts)

        except Exception as e:
            logger.error(f"DOCX parsing failed: {e}")
            raise ValueError(f"Failed to parse DOCX: {e}")

    async def parse_cv(
        self,
        cv: CV,
        file_content: bytes,
        db: Session,
        user_id: str,
    ) -> Dict[str, Any]:
        """
        Parse CV file and extract structured data using LLM

        Args:
            cv: CV database object
            file_content: Raw file bytes
            db: Database session
            user_id: User ID for tracking

        Returns:
            Parsed CV data dictionary
        """
        try:
            # Extract text based on file extension
            filename_lower = cv.filename.lower()

            if filename_lower.endswith('.pdf'):
                cv_text = self.extract_text_from_pdf(file_content)
            elif filename_lower.endswith('.docx') or filename_lower.endswith('.doc'):
                cv_text = self.extract_text_from_docx(file_content)
            else:
                raise ValueError(f"Unsupported file format: {cv.filename}")

            if not cv_text or len(cv_text.strip()) < 100:
                raise ValueError("Extracted text is too short or empty")

            # Save parsed text to CV
            cv.parsed_text = cv_text
            db.commit()

            # Prepare prompt
            prompt = CV_PARSING_PROMPT.format(cv_text=cv_text)

            # Call LLM to parse CV
            result = await bedrock_service.invoke_claude(
                prompt=prompt,
                db=db,
                user_id=user_id,
                call_type=LLMCallType.CV_PARSING,
                cv_id=str(cv.id),
                max_tokens=8000,  # Need more tokens for detailed CV parsing
                temperature=0.3,  # Lower temperature for more consistent extraction
            )

            if not result["success"]:
                raise ValueError(f"LLM parsing failed: {result.get('error')}")

            # Parse JSON response
            try:
                parsed_data = json.loads(result["response"])
            except json.JSONDecodeError as e:
                logger.error(f"Failed to parse LLM response as JSON: {e}")
                logger.error(f"Response: {result['response'][:500]}")
                raise ValueError(f"Invalid JSON response from LLM: {e}")

            # Extract quick access fields
            personal_info = parsed_data.get("personal_info", {})
            summary = parsed_data.get("summary", {})
            meta = parsed_data.get("meta", {})
            skill_freshness = meta.get("skill_freshness_summary", {})

            # Create CVParseDetail record
            cv_parse_detail = CVParseDetail(
                cv_id=cv.id,
                user_id=user_id,
                parsed_data=parsed_data,
                candidate_name=personal_info.get("name"),
                candidate_email=personal_info.get("email"),
                current_role=summary.get("current_role"),
                current_company=summary.get("current_company"),
                total_experience_years=summary.get("total_experience_years"),
                career_level=summary.get("career_level"),
                current_skills_count=skill_freshness.get("current_skills_count", 0),
                outdated_skills_count=skill_freshness.get("outdated_skills_count", 0),
                github_username=personal_info.get("github", "").split("/")[-1] if personal_info.get("github") else None,
                cv_quality_score=self._calculate_quality_score(meta.get("cv_quality", "medium")),
                parsing_confidence=meta.get("parsing_confidence"),
                red_flags_count=len(meta.get("red_flags", [])),
            )

            db.add(cv_parse_detail)
            db.commit()
            db.refresh(cv_parse_detail)

            logger.info(f"Successfully parsed CV {cv.id} for candidate: {cv_parse_detail.candidate_name}")

            return {
                "success": True,
                "parse_detail_id": str(cv_parse_detail.id),
                "parsed_data": parsed_data,
                "usage": result["usage"],
                "cost": result["cost"],
            }

        except Exception as e:
            logger.error(f"CV parsing error: {e}")
            return {
                "success": False,
                "error": str(e),
            }

    def _calculate_quality_score(self, quality: str) -> int:
        """Convert quality string to numeric score"""
        quality_map = {
            "high": 85,
            "medium": 60,
            "low": 35,
        }
        return quality_map.get(quality, 50)


# Singleton instance
cv_parser_service = CVParserService()
